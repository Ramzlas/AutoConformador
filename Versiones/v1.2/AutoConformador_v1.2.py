import csv
from docx import Document
from docx.shared import Pt

nombre_archivo = input("Ingrese el nombre del archivo CSV (no incluya .csv): ")  # sin extensión
archivo_csv = nombre_archivo + ".csv"
modelo_docx = "modelo.docx"  # Documento base

columnas_clave = {
    "factura": ["nº factura o tasa"],
    "monto": ["monto"],
    "recurrente": ["recurrente"],
    "suministro": ["nº suministro"]
}

def encontrar_columna(candidatas, encabezados):
    for col in candidatas:
        for encabezado in encabezados:
            if col.strip().lower() == encabezado.strip().lower():
                return encabezado
    return None

try:
    with open(archivo_csv, newline='', encoding='utf-8-sig') as csvfile:
        lector = csv.DictReader(csvfile, delimiter=';')
        encabezados = lector.fieldnames

        col_factura = encontrar_columna(columnas_clave["factura"], encabezados)
        col_monto = encontrar_columna(columnas_clave["monto"], encabezados)
        col_recurrente = encontrar_columna(columnas_clave["recurrente"], encabezados)
        col_suministro = encontrar_columna(columnas_clave["suministro"], encabezados)

        if None in [col_factura, col_monto, col_recurrente, col_suministro]:
            raise KeyError("No se encontraron todas las columnas necesarias. Verifique los nombres.")

        for fila in lector:
            factura = fila[col_factura].strip()
            monto_raw = fila[col_monto].strip()
            proveedor = fila[col_recurrente].strip()
            suministro = fila[col_suministro].strip()

            # Limpieza
            monto_neto = monto_raw.replace("$", "").replace(".", "").replace(",", ".")
            try:
                monto = float(monto_neto)
            except ValueError:
                print(f"⚠️ Error con el monto: {monto_raw} (salteando fila)")
                continue

            es_tasa = factura.lower().startswith("tasa")
            factura_limpia = factura.replace("Factura Nº", "").replace("Tasa Nº", "").strip()
            suministro_limpio = suministro.replace("Nº", "").strip()
            monto_formateado = f"{monto:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

            # Cargar modelo existente
            doc = Document(modelo_docx)

            from docx.oxml.ns import qn  # necesario para fuente en español

            for p in doc.paragraphs:
                if "Espacio para el cuerpo del texto" in p.text:
                    # Borrar el contenido original del párrafo
                    p.clear()

                    def agregar_run(texto, bold=False):
                        run = p.add_run(texto)
                        run.bold = bold
                        run.italic = False  # asegurarse que no esté en itálica
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        run.font.size = Pt(12)
                        return run

                    agregar_run("                                                            Visto el informe que antecede, esta Dirección presta conformidad a la ")
                    agregar_run("tasa" if es_tasa else "factura")
                    agregar_run(f" Nº{factura_limpia}", bold=True)
                    agregar_run(" por ")
                    agregar_run(f"$ {monto_formateado}", bold=True) 
                    agregar_run(" de ")
                    agregar_run(proveedor, bold=True)
                    agregar_run(" referente al suministro ")
                    agregar_run(f"Nº{suministro_limpio}", bold=True)
                    agregar_run(" para su liquidación y pago.")
                    break


            # Guardar el archivo con nombre personalizado
            nombre_archivo_salida = f"Nota_suministro_{suministro_limpio}.docx"
            doc.save(nombre_archivo_salida)
            print(f"✅ Generado: {nombre_archivo_salida}")

except FileNotFoundError:
    print("❌ El archivo no se encontró:", archivo_csv)
except KeyError as e:
    print(f"❌ Faltan columnas esperadas en el archivo: {e}")
except Exception as e:
    print(f"❌ Ocurrió un error inesperado: {e}")
