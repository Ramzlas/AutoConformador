import csv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Solicita el nombre del archivo CSV a procesar
nombre_archivo = input("Ingrese el nombre del archivo CSV (no incluya .csv): ")
archivo_csv = nombre_archivo + ".csv"
archivo_salida = "Listado.docx"

# Define nombres posibles de columnas clave
columnas_clave = {
    "factura": ["nº factura o tasa"],
    "monto": ["monto"],
    "recurrente": ["recurrente"],
    "suministro": ["nº suministro"]
}

# Busca la columna correspondiente según variantes posibles
def encontrar_columna(candidatas, encabezados):
    for col in candidatas:
        for encabezado in encabezados:
            if col.strip().lower() == encabezado.strip().lower():
                return encabezado
    return None

try:
    # Abre y lee el archivo CSV
    with open(archivo_csv, newline='', encoding='utf-8-sig') as csvfile:
        lector = csv.DictReader(csvfile, delimiter=';')
        encabezados = lector.fieldnames

        # Identifica las columnas necesarias
        col_factura = encontrar_columna(columnas_clave["factura"], encabezados)
        col_monto = encontrar_columna(columnas_clave["monto"], encabezados)
        col_recurrente = encontrar_columna(columnas_clave["recurrente"], encabezados)
        col_suministro = encontrar_columna(columnas_clave["suministro"], encabezados)

        if None in [col_factura, col_monto, col_recurrente, col_suministro]:
            raise KeyError("No se encontraron todas las columnas necesarias. Verifique los nombres.")

        # Crea documento .DOCX
        doc = Document()

        # Cambia la fuente y tamaño por defecto del estilo Normal para todo el documento
        estilo_normal = doc.styles['Normal']
        estilo_normal.font.name = 'Times New Roman'
        estilo_normal.font.size = Pt(12)

        # Para que la fuente se aplique correctamente en Word:
        from docx.oxml.ns import qn
        estilo_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        for fila in lector:
            # Extrae y limpia datos por fila
            factura = fila[col_factura].strip()
            monto_raw = fila[col_monto].strip()
            proveedor = fila[col_recurrente].strip()
            suministro = fila[col_suministro].strip()

            # Convierte el monto a número
            monto_neto = monto_raw.replace("$", "").replace(".", "").replace(",", ".")
            try:
                monto = float(monto_neto)
            except ValueError:
                print(f"⚠️ Error con el monto: {monto_raw} (salteando fila)")
                continue

            # Define tipo de documento y limpia texto
            es_tasa = factura.lower().startswith("tasa")
            factura_limpia = factura.replace("Factura Nº", "").replace("Tasa Nº", "").strip()
            suministro_limpio = suministro.replace("Nº", "").strip()

            # Agrega texto al documento
            p = doc.add_paragraph()
            inicio_texto = "Visto el informe que antecede, esta Secretaría presta conformidad a la "
            tipo_documento = "tasa" if es_tasa else "factura"

            p.add_run(inicio_texto)
            p.add_run(tipo_documento)
            p.add_run(" Nº" + factura_limpia).bold = True
            p.add_run(" por ")
            monto_formateado = f"{monto:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            p.add_run("$ " + monto_formateado).bold = True
            p.add_run(" de ")
            p.add_run(proveedor).bold = True
            p.add_run(" referente al suministro ")
            p.add_run("Nº" + suministro_limpio).bold = True
            p.add_run(" para su liquidación y pago.")

            doc.add_paragraph()  # Espacio entre párrafos
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Guarda el documento final
        doc.save(archivo_salida)
        print(f"✅ Documento generado: '{archivo_salida}'")

# Manejo de errores
except FileNotFoundError:
    print("❌ El archivo no se encontró:", archivo_csv)
except KeyError as e:
    print(f"❌ Faltan columnas esperadas en el archivo: {e}")
except Exception as e:
    print(f"❌ Ocurrió un error inesperado: {e}")
