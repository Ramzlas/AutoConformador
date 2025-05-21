import csv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Pide nombre del archivo CSV (sin extensión)
nombre_archivo = input("Ingrese el nombre del archivo CSV (no incluya .csv): ")
archivo_csv = nombre_archivo + ".csv"

# Lista de posibles nombres de columnas para mayor flexibilidad
columnas_clave = {
    "factura": ["nº factura o tasa", "factura", "número de factura", "tasa"],
    "monto": ["monto", "importe"],
    "recurrente": ["recurrente", "proveedor", "razón social"],
    "suministro": ["nº suministro", "suministro"]
}

# Busca la columna que coincida entre las opciones
def encontrar_columna(candidatas, encabezados):
    for col in candidatas:
        for encabezado in encabezados:
            if col.strip().lower() == encabezado.strip().lower():
                return encabezado
    return None

try:
    # Abre y procesa el archivo CSV
    with open(archivo_csv, newline='', encoding='utf-8-sig') as csvfile:
        lector = csv.DictReader(csvfile, delimiter=';')
        encabezados = lector.fieldnames

        # Determina columnas importantes
        col_factura = encontrar_columna(columnas_clave["factura"], encabezados)
        col_monto = encontrar_columna(columnas_clave["monto"], encabezados)
        col_recurrente = encontrar_columna(columnas_clave["recurrente"], encabezados)
        col_suministro = encontrar_columna(columnas_clave["suministro"], encabezados)

        if None in [col_factura, col_monto, col_recurrente, col_suministro]:
            raise KeyError("No se encontraron todas las columnas necesarias. Verifique los nombres.")

        for fila in lector:
            # Limpieza y extracción de valores
            factura = fila[col_factura].strip()
            monto_raw = fila[col_monto].strip()
            proveedor = fila[col_recurrente].strip()
            suministro = fila[col_suministro].strip()

            monto_neto = monto_raw.replace("$", "").replace(".", "").replace(",", ".")
            try:
                monto = float(monto_neto)
            except ValueError:
                print(f"⚠️ Error con el monto: {monto_raw} (salteando fila)")
                continue

            es_tasa = factura.lower().startswith("tasa")
            factura_limpia = factura.replace("Factura Nº", "").replace("Tasa Nº", "").strip()
            suministro_limpio = suministro.replace("Nº", "").strip()
            monto_formateado = f"{monto:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

            # Crea documento por cada fila
            doc = Document()

            # Configura estilo del documento
            style = doc.styles['Normal']
            style.font.name = 'Times new roman'
            style.font.size = Pt(12)

            # Encabezado institucional
            def agregar_parrafo_encabezado(doc, texto, subrayado=False):
                p = doc.add_paragraph()
                run = p.add_run(texto)
                run.underline = subrayado
                p.paragraph_format.space_after = Pt(0)

            agregar_parrafo_encabezado(doc, "Secretaría de Coordinación Administrativa")
            agregar_parrafo_encabezado(doc, "Sr. Alejandro Fernández")
            agregar_parrafo_encabezado(doc, "S                /                D.", subrayado=True)
            doc.add_paragraph()

            # Cuerpo del documento
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            inicio_texto = "                                                            Visto el informe que antecede, esta Secretaría presta conformidad a la "
            tipo_documento = "tasa" if es_tasa else "factura"

            p.add_run(inicio_texto)
            p.add_run(tipo_documento)
            p.add_run(" Nº" + factura_limpia).bold = True
            p.add_run(" por ")
            p.add_run("$ " + monto_formateado).bold = True
            p.add_run(" de ")
            p.add_run(proveedor).bold = True
            p.add_run(" referente al suministro ")
            p.add_run("Nº" + suministro_limpio).bold = True
            p.add_run(" para su liquidación y pago.")

            doc.add_paragraph()
            doc.add_paragraph("Sin otro particular, saludo atentamente.")

            # Guarda cada nota generada
            nombre_archivo_salida = f"Nota_suministro_{suministro_limpio}.docx"
            doc.save(nombre_archivo_salida)
            print(f"✅ Generado: {nombre_archivo_salida}")

except FileNotFoundError:
    print("❌ El archivo no se encontró:", archivo_csv)
except KeyError as e:
    print(f"❌ Faltan columnas esperadas en el archivo: {e}")
except Exception as e:
    print(f"❌ Ocurrió un error inesperado: {e}")